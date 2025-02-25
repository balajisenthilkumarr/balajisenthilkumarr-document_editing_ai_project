import json
import os
import random
import re
from docx import Document

# Define transformation commands and ensure text contains relevant content
COMMANDS = [
    {"command": "Remove bullet points", "required": "•"},
    {"command": "Change 'January' to 'February'", "required": "January"},
    {"command": "Delete all occurrences of 'confidential'", "required": "confidential"},
    {"command": "Replace 'project' with 'assignment'", "required": "project"},
    {"command": "Remove numbers", "required": "[0-9]"},
    {"command": "Replace 'CEO' with 'Managing Director'", "required": "CEO"},
    {"command": "Convert to uppercase", "required": "[a-z]"},
    {"command": "Remove the first paragraph", "required": "\n\n"},
    {"command": "Capitalize all words", "required": "[a-z]"},
    {"command": "Remove the last sentence", "required": "\\."}, # Fixed escape sequence
    {"command": "Change all dates to next year", "required": "202[0-9]"},
    {"command": "Replace 'budget' with 'financial plan'", "required": "budget"},
    {"command": "Remove all percentage figures", "required": "%"},
    {"command": "Convert all department names to lowercase", "required": "HR|IT|Marketing|Sales|Finance"},
    {"command": "Replace 'team' with 'group'", "required": "team"},
    {"command": "Add 'DRAFT: ' to the beginning of each paragraph", "required": ".+"},
    {"command": "Remove all occurrences of 'quarterly'", "required": "quarterly"},
    {"command": "Change all instances of 'review' to 'assessment'", "required": "review"}
]

def generate_content_for_command(command_info):
    """Generate content that is suitable for the specific command."""
    command = command_info["command"]
    required = command_info["required"]
    
    # Base paragraphs with varied content
    all_paragraphs = [
        "The CEO will address the employees regarding company growth in January.",
        "The budget review is planned for the upcoming quarterly meeting.",
        "Upcoming deadlines include March 5, 2024, and July 20, 2025.",
        "Customer feedback indicates satisfaction level of 85%.",
        "Department HR requested additional resources for Q3.",
        "International offices reported strong growth this quarter.",
        "The innovation team proposed AI-powered analytics as a new project.",
        "• First quarter results exceeded expectations by 12%.",
        "• The Marketing team achieved 95% of their quarterly goals.",
        "Our confidential report shows significant improvements in Q2.",
        "The Sales team's project is scheduled for review on January 15.",
        "The IT department submitted their quarterly budget of $250,000."
    ]
    
    # For special cases that need specific content
    if "Remove the first paragraph" in command:
        return "\n\n".join(random.sample(all_paragraphs, k=random.randint(2, 4)))
    
    elif "Remove the last sentence" in command:
        multi_sentence_paragraphs = [
            "The CEO presented the new strategy. Teams will implement it gradually. Results are expected by Q3.",
            "Our quarterly review is complete. The results were positive. We exceeded expectations in most areas.",
            "The budget allocation is finalized. Each department received their requested funds. Implementation begins next month."
        ]
        selected = random.choice(multi_sentence_paragraphs)
        return selected + "\n\n" + "\n\n".join(random.sample(all_paragraphs, k=random.randint(1, 3)))

    # Filter paragraphs that match the required pattern
    paragraphs = [para for para in all_paragraphs if re.search(required, para, re.IGNORECASE)]

    # Handle empty case by adding fallback content
    if not paragraphs:
        if "bullet points" in command:
            paragraphs = ["• First item in the list\n• Second important item\n• Final consideration"]
        elif "January" in command:
            paragraphs = ["The board meeting is scheduled for January 15th."]
        elif "confidential" in command:
            paragraphs = ["Please treat this information as confidential until further notice."]
        elif "project" in command:
            paragraphs = ["The new project timeline has been approved by management."]
        elif "CEO" in command:
            paragraphs = ["The CEO has requested a full audit of all departments."]
        elif "budget" in command:
            paragraphs = ["The annual budget must be submitted by the end of Q4."]
        elif "percentage" in command or "%" in command:
            paragraphs = ["Our efficiency improved by 45% after implementing the new system."]
        elif "team" in command:
            paragraphs = ["The development team will present their findings next week."]
        elif "quarterly" in command:
            paragraphs = ["The quarterly report shows positive trends in all markets."]
        elif "review" in command:
            paragraphs = ["The performance review process will begin next month."]
        else:
            paragraphs = ["The company is making progress on multiple fronts in Q1 of 2024."]

    # Select extra paragraphs without exceeding available ones
    available_extra = [p for p in all_paragraphs if p not in paragraphs]
    k = min(len(available_extra), random.randint(1, 3))  # Avoid out-of-range errors
    extra_paragraphs = random.sample(available_extra, k) if available_extra else []
    
    all_selected = paragraphs + extra_paragraphs
    random.shuffle(all_selected)

    return "\n\n".join(all_selected)


def transform_content(content, command):
    """Apply transformation based on command."""
    if command == "Capitalize all words":
        return ' '.join(word.capitalize() for word in content.split())

    elif command == "Replace 'CEO' with 'Managing Director'":
        return content.replace("CEO", "Managing Director")

    elif command == "Convert to uppercase":
        return content.upper()

    elif command == "Delete all occurrences of 'confidential'":
        return content.replace("confidential", "").replace("Confidential", "").strip()

    elif command == "Remove numbers":
        return ''.join(char for char in content if not char.isdigit())

    elif command == "Change 'January' to 'February'":
        return content.replace("January", "February").replace("january", "february")

    elif command == "Remove bullet points":
        return content.replace("• ", "").replace("•", "")

    elif command == "Remove the first paragraph":
        paragraphs = content.split('\n\n')
        return '\n\n'.join(paragraphs[1:]) if len(paragraphs) > 1 else ""

    elif command == "Replace 'project' with 'assignment'":
        return content.replace("project", "assignment").replace("Project", "Assignment")

    elif command == "Remove the last sentence":
        content_parts = content.split('\n\n')
        result_parts = []
        
        for part in content_parts:
            sentences = re.split(r'(?<=[.!?])\s+', part.strip())
            if len(sentences) > 1:
                result_parts.append(' '.join(sentences[:-1]))
            else:
                result_parts.append(part)
        
        return '\n\n'.join(result_parts)
    
    elif command == "Change all dates to next year":
        def increment_year(match):
            year = int(match.group(0))
            return str(year + 1)
        return re.sub(r'202\d', increment_year, content)
    
    elif command == "Replace 'budget' with 'financial plan'":
        return content.replace("budget", "financial plan").replace("Budget", "Financial plan")
    
    elif command == "Remove all percentage figures":
        return re.sub(r'\d+%', '', content)
    
    elif command == "Convert all department names to lowercase":
        for dept in ["HR", "IT", "Marketing", "Sales", "Finance"]:
            content = content.replace(dept, dept.lower())
        return content
    
    elif command == "Replace 'team' with 'group'":
        return content.replace("team", "group").replace("Team", "Group")
    
    elif command == "Add 'DRAFT: ' to the beginning of each paragraph":
        paragraphs = content.split('\n\n')
        return '\n\n'.join([f"DRAFT: {p}" for p in paragraphs])
    
    elif command == "Remove all occurrences of 'quarterly'":
        return content.replace("quarterly", "").replace("Quarterly", "")
    
    elif command == "Change all instances of 'review' to 'assessment'":
        return content.replace("review", "assessment").replace("Review", "Assessment")
    
    return content

def generate_dataset(output_dir, num_documents=1500):
    """Generate dataset with original and modified documents and save to disk."""
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)

    metadata = []

    for i in range(1, num_documents + 1):
        # Select a command
        command_info = random.choice(COMMANDS)
        command = command_info["command"]
        
        # Generate appropriate content for this command
        original_content = generate_content_for_command(command_info)
        
        # Apply the transformation
        modified_content = transform_content(original_content, command)
        
        # Ensure transformation actually changed something
        attempt = 0
        while original_content == modified_content and attempt < 3:
            # Try a different content if transformation didn't work
            original_content = generate_content_for_command(command_info)
            modified_content = transform_content(original_content, command)
            attempt += 1
            
        # If still no change after attempts, skip this document
        if original_content == modified_content:
            print(f"⚠️ Warning: Couldn't apply '{command}' effectively. Skipping document {i}.")
            continue

        doc_entry = {
            "document_id": f"doc_{i}",
            "original_file": f"doc_{i}_original.docx",
            "modified_file": f"doc_{i}_modified.docx",
            "command": command
        }

        # Save original document
        original_doc = Document()
        for paragraph in original_content.split('\n\n'):
            original_doc.add_paragraph(paragraph)
        original_path = os.path.join(output_dir, doc_entry["original_file"])
        original_doc.save(original_path)

        # Save modified document
        modified_doc = Document()
        for paragraph in modified_content.split('\n\n'):
            modified_doc.add_paragraph(paragraph)
        modified_path = os.path.join(output_dir, doc_entry["modified_file"])
        modified_doc.save(modified_path)

        metadata.append(doc_entry)

        if i % 100 == 0:
            print(f"✅ {i} documents generated...")

    metadata_path = os.path.join(output_dir, 'metadata.json')
    with open(metadata_path, 'w', encoding='utf-8') as f:
        json.dump(metadata, f, indent=4)

    return metadata_path

if __name__ == "__main__":
    output_directory = "data/training_dataset"
    metadata_file = generate_dataset(output_directory, num_documents=1500)
    print(f"\n🎉 Dataset generated successfully!")
    print(f"📄 Metadata file: {metadata_file}")